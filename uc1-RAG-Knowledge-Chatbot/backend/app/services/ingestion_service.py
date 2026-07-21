import asyncio
import hashlib
import logging
import re
from dataclasses import dataclass
from pathlib import Path

import tiktoken
from azure.search.documents.indexes.models import (
    HnswAlgorithmConfiguration,
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    SemanticConfiguration,
    SemanticField,
    SemanticPrioritizedFields,
    SemanticSearch,
    SimpleField,
    VectorSearch,
    VectorSearchProfile,
)
from azure.search.documents.aio import SearchClient
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from openai import AsyncOpenAI
from pypdf import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession

from app.clients.azure_foundry_client import get_openai_client
from app.clients.azure_search_client import get_search_client, get_search_index_client
from app.config import get_settings
from app.db import async_session_factory
from app.repositories import document_repo, indexing_run_repo
from app.services.generation_service import embed_text

logger = logging.getLogger(__name__)
settings = get_settings()

RESOURCES_DIR = Path(__file__).resolve().parents[2] / "resources"

CHUNK_MAX_TOKENS = 700
CHUNK_OVERLAP_RATIO = 0.12

_encoding = tiktoken.get_encoding("cl100k_base")

@dataclass
class LoadedDocument:
    format: str
    text: str  # normalized text; heading lines prefixed with '#' * level


@dataclass
class Chunk:
    chunk_ref: str
    content: str


_HTML_HEAD_RE = re.compile(rb"^\s*(<!doctype\s+html|<html[\s>])", re.IGNORECASE)


def sniff_format(path: Path) -> str:
    """Sniffs actual content type from the file's bytes rather than trusting the
    extension, in case a future corpus file is mislabeled the way a couple of earlier
    ones were. Previously used python-magic, but that needs the system libmagic1
    library, which isn't available (and isn't reliably installable) in Azure App
    Service's Python code-deployment runtime -- replaced with a small dependency-free
    check covering exactly the formats this corpus actually has."""
    ext = path.suffix.lower().lstrip(".")
    with path.open("rb") as f:
        head = f.read(4096)

    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK\x03\x04"):
        return "docx"
    if _HTML_HEAD_RE.match(head):
        return "html"
    # No binary signature matched -- it's text of some kind. There's no content-level
    # way to tell markdown from plain text (both are just prose), so the extension is
    # the only signal for that specific distinction.
    if ext in ("md", "markdown"):
        return "markdown"
    return ext or "unknown"


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"# Page {i}\n{text.strip()}")
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = ((para.style.name if para.style else None) or "").lower()
        match = re.match(r"heading (\d)", style)
        if match:
            level = min(int(match.group(1)), 6)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)
    return "\n\n".join(lines)


def _load_html(path: Path) -> str:
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    lines = []
    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        text = el.get_text(strip=True)
        if not text:
            continue
        if el.name.startswith("h"):
            level = int(el.name[1])
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)
    return "\n\n".join(lines)


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


_LOADERS = {
    "pdf": _load_pdf,
    "docx": _load_docx,
    "html": _load_html,
    "text": _load_text,
    "markdown": _load_text,
}


def load_document(path: Path) -> LoadedDocument:
    fmt = sniff_format(path)
    loader = _LOADERS.get(fmt)
    if loader is None:
        raise ValueError(f"Unsupported format '{fmt}' for {path.name}")
    return LoadedDocument(format=fmt, text=loader(path))


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _split_sections(text: str) -> list[tuple[str, str]]:
    """Splits normalized text into (heading, body) sections. Text before the first
    heading (or all of it, if there are no headings) becomes one section titled 'Document'."""
    sections: list[tuple[str, str]] = []
    current_heading = "Document"
    current_lines: list[str] = []

    for line in text.splitlines():
        match = _HEADING_RE.match(line)
        if match:
            if current_lines:
                sections.append((current_heading, "\n".join(current_lines).strip()))
            current_heading = match.group(2).strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_heading, "\n".join(current_lines).strip()))
    return [(h, b) for h, b in sections if b]


def _split_by_tokens(text: str, max_tokens: int, overlap_ratio: float) -> list[str]:
    tokens = _encoding.encode(text)
    if len(tokens) <= max_tokens:
        return [text]
    step = max(1, int(max_tokens * (1 - overlap_ratio)))
    pieces = []
    for start in range(0, len(tokens), step):
        piece_tokens = tokens[start : start + max_tokens]
        pieces.append(_encoding.decode(piece_tokens))
        if start + max_tokens >= len(tokens):
            break
    return pieces


def chunk_document(loaded: LoadedDocument) -> list[Chunk]:
    chunks: list[Chunk] = []
    sections = _split_sections(loaded.text)
    for heading, body in sections:
        pieces = _split_by_tokens(body, CHUNK_MAX_TOKENS, CHUNK_OVERLAP_RATIO)
        for i, piece in enumerate(pieces, start=1):
            chunk_ref = heading if len(pieces) == 1 else f"{heading} (part {i}/{len(pieces)})"
            chunks.append(Chunk(chunk_ref=chunk_ref, content=piece))
    return chunks


def _checksum(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


SEMANTIC_CONFIG_NAME = "uc1-semantic-config"


async def ensure_index_exists() -> None:
    async with get_search_index_client() as index_client:
        existing = [name async for name in index_client.list_index_names()]
        if settings.azure_search_index_name in existing:
            return

        index = SearchIndex(
            name=settings.azure_search_index_name,
            fields=[
                SimpleField(name="id", type=SearchFieldDataType.String, key=True),
                SimpleField(name="document_id", type=SearchFieldDataType.String, filterable=True),
                SearchableField(name="filename", type=SearchFieldDataType.String, filterable=True),
                SearchableField(name="chunk_ref", type=SearchFieldDataType.String),
                SearchableField(name="content", type=SearchFieldDataType.String),
                SearchField(
                    name="content_vector",
                    type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                    searchable=True,
                    vector_search_dimensions=3072,
                    vector_search_profile_name="uc1-vector-profile",
                ),
            ],
            vector_search=VectorSearch(
                algorithms=[HnswAlgorithmConfiguration(name="uc1-hnsw")],
                profiles=[
                    VectorSearchProfile(name="uc1-vector-profile", algorithm_configuration_name="uc1-hnsw")
                ],
            ),
            semantic_search=SemanticSearch(
                configurations=[
                    SemanticConfiguration(
                        name=SEMANTIC_CONFIG_NAME,
                        prioritized_fields=SemanticPrioritizedFields(
                            title_field=SemanticField(field_name="chunk_ref"),
                            content_fields=[SemanticField(field_name="content")],
                        ),
                    )
                ]
            ),
        )
        await index_client.create_index(index)


EMBED_CONCURRENCY = 20

# Documents used to ingest strictly one at a time -- correct for durability (each one
# commits only after its own Search upload succeeds, see the per-document comment below)
# but meant total wall-clock time was the *sum* of every document's time, not the max.
# This corpus is lopsided (hr-manual.docx alone is 132 of 481 chunks), so the four
# largest documents dominated a ~90s run despite chunk-level embedding already being
# concurrent. Running documents concurrently too (bounded, so as not to spike Azure
# Search/Foundry traffic all at once) cuts wall-clock roughly toward the slowest single
# document instead of the sum of all of them, without weakening the per-document
# atomicity -- each document still gets its own dedicated DB session and commits
# independently, since AsyncSession isn't safe to share across concurrent coroutines.
DOC_CONCURRENCY = 4


async def _ingest_one(
    path: Path,
    search_client: SearchClient,
    embed_client: AsyncOpenAI,
    embed_semaphore: asyncio.Semaphore,
) -> int:
    async def embed_bounded(text: str) -> list[float]:
        async with embed_semaphore:
            return await embed_text(text, client=embed_client)

    # PDF/DOCX/HTML parsing and tiktoken tokenization are synchronous, CPU-bound work --
    # running them directly on the event loop would freeze every other request (chat,
    # document list, etc.) on the whole process for as long as they take. asyncio.to_thread
    # hands each off to a worker thread so the event loop stays free to serve other traffic
    # while this runs.
    checksum = await asyncio.to_thread(_checksum, path)
    fmt = await asyncio.to_thread(sniff_format, path)
    loaded = await asyncio.to_thread(load_document, path)
    chunks = await asyncio.to_thread(chunk_document, loaded)

    async with async_session_factory() as doc_db:
        document = await document_repo.upsert_document(
            doc_db, filename=path.name, format=fmt, source_path=str(path), checksum=checksum
        )

        vectors = await asyncio.gather(*(embed_bounded(chunk.content) for chunk in chunks))
        search_docs = [
            {
                "id": f"{document.id}-{i}",
                "document_id": str(document.id),
                "filename": document.filename,
                "chunk_ref": chunk.chunk_ref,
                "content": chunk.content,
                "content_vector": vector,
            }
            for i, (chunk, vector) in enumerate(zip(chunks, vectors, strict=True))
        ]

        if search_docs:
            await search_client.upload_documents(search_docs)

        await document_repo.mark_indexed(doc_db, document.id, len(chunks))
        await doc_db.commit()  # durable per-document: a later failure (this document or
        # another running concurrently) can't orphan this document's Search entries
        # against a rolled-back Postgres row (see incident in commit history)

    logger.info("ingested %s: %d chunks", path.name, len(chunks))
    return len(chunks)


async def _prune_removed_documents(search_client: SearchClient) -> int:
    """Marks documents whose source file no longer exists in RESOURCES_DIR as removed and
    deletes their chunks from Azure Search. A plain reindex only upserts files it finds on
    disk -- it never notices a file that was deleted from the corpus, so without this a
    removed file's document row (and its now-orphaned Search chunks) would linger forever,
    showing up in the admin document list and remaining retrievable/citable indefinitely."""
    current_filenames = {path.name for path in RESOURCES_DIR.iterdir() if path.is_file()}

    async with async_session_factory() as db:
        stale = await document_repo.list_stale_documents(db, current_filenames)
        for document in stale:
            results = await search_client.search(
                search_text="*", filter=f"document_id eq '{document.id}'", select=["id"], top=1000
            )
            stale_ids = [{"id": r["id"]} async for r in results]
            if stale_ids:
                await search_client.delete_documents(stale_ids)
            await document_repo.mark_removed(db, document.id)
        await db.commit()

    return len(stale)


async def _process_corpus() -> tuple[int, int]:
    """Parses, chunks, embeds, and indexes every file in RESOURCES_DIR. Returns
    (doc_count, chunk_count). Shared by the synchronous CLI path and the background-task
    path -- neither owns a caller-supplied db session, since ingesting each document
    already opens its own (see _ingest_one)."""
    await ensure_index_exists()

    embed_semaphore = asyncio.Semaphore(EMBED_CONCURRENCY)
    doc_semaphore = asyncio.Semaphore(DOC_CONCURRENCY)

    async def ingest_one_bounded(path: Path, search_client: SearchClient, embed_client: AsyncOpenAI) -> int:
        async with doc_semaphore:
            return await _ingest_one(path, search_client, embed_client, embed_semaphore)

    async with get_search_client() as search_client, get_openai_client() as embed_client:
        paths = [path for path in sorted(RESOURCES_DIR.iterdir()) if path.is_file()]
        chunk_counts = await asyncio.gather(*(ingest_one_bounded(path, search_client, embed_client) for path in paths))
        await _prune_removed_documents(search_client)

    return len(paths), sum(chunk_counts)


async def ingest_all(db: AsyncSession, triggered_by: str = "admin") -> dict:
    run = await indexing_run_repo.start_run(db, triggered_by=triggered_by)
    total_docs, total_chunks = await _process_corpus()

    await indexing_run_repo.finish_run(
        db, run.id, doc_count=total_docs, chunk_count=total_chunks, status="completed"
    )
    await db.commit()
    return {"doc_count": total_docs, "chunk_count": total_chunks}


async def run_ingestion_background(run_id) -> None:
    """Entry point for the reindex endpoint's BackgroundTasks call: the endpoint has
    already created and returned the IndexingRun row on its own request-scoped session
    before scheduling this, so this runs entirely on its own fresh session (the request's
    session is closed by the time a background task executes) and only has to update that
    same row when the corpus finishes processing -- or mark it failed, so a crash here
    doesn't leave the row stuck at "running" forever with no visible error."""
    try:
        total_docs, total_chunks = await _process_corpus()
        status = "completed"
    except Exception:
        logger.exception("ingestion_failed", extra={"run_id": str(run_id)})
        total_docs, total_chunks = 0, 0
        status = "failed"

    async with async_session_factory() as db:
        await indexing_run_repo.finish_run(db, run_id, doc_count=total_docs, chunk_count=total_chunks, status=status)
        await db.commit()
