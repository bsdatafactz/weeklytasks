from pathlib import Path

from docx import Document as DocxDocument

from app.services.ingestion_service import (
    CHUNK_MAX_TOKENS,
    LoadedDocument,
    _encoding,
    _split_by_tokens,
    _split_sections,
    chunk_document,
    load_document,
    sniff_format,
)

RESOURCES_DIR = Path(__file__).resolve().parents[1] / "resources"


class TestSniffFormat:
    def test_misleading_html_extensions_are_sniffed_as_html(self):
        # These two corpus files have non-HTML extensions but are actually HTML content --
        # the whole point of content-sniffing instead of trusting the extension.
        assert sniff_format(RESOURCES_DIR / "progressive-discipline-policy") == "html"
        assert sniff_format(RESOURCES_DIR / "employee-handbook-sample.doc") == "html"

    def test_real_pdf_and_docx_still_sniff_correctly(self):
        assert sniff_format(RESOURCES_DIR / "attendance-policy.pdf") == "pdf"
        assert sniff_format(RESOURCES_DIR / "remote-work-policy.docx") == "docx"

    def test_markdown_extension_disambiguated_from_plain_text(self):
        # `magic` reports .md as generic text/plain; the extension is the only signal
        # that distinguishes markdown from plain text once content-sniffing says text/plain.
        assert sniff_format(RESOURCES_DIR / "benefits-and-perks.md") == "markdown"


class TestSplitSections:
    def test_text_before_first_heading_becomes_document_section(self):
        text = "intro line\n\n# Heading One\nbody one"
        sections = _split_sections(text)
        assert sections[0] == ("Document", "intro line")
        assert sections[1] == ("Heading One", "body one")

    def test_no_headings_at_all_is_one_document_section(self):
        text = "just plain body text\nwith no structure"
        sections = _split_sections(text)
        assert len(sections) == 1
        assert sections[0][0] == "Document"

    def test_empty_sections_are_dropped(self):
        text = "# Heading A\n\n# Heading B\nsome body"
        sections = _split_sections(text)
        # Heading A has no body before Heading B starts -- should not produce an empty chunk.
        assert sections == [("Heading B", "some body")]


class TestSplitByTokens:
    def test_short_text_is_not_split(self):
        text = "a few words, well under the token budget"
        pieces = _split_by_tokens(text, CHUNK_MAX_TOKENS, 0.12)
        assert pieces == [text]

    def test_long_text_splits_into_multiple_overlapping_pieces(self):
        text = " ".join(f"word{i}" for i in range(2000))
        pieces = _split_by_tokens(text, max_tokens=100, overlap_ratio=0.12)
        assert len(pieces) > 1
        for piece in pieces:
            assert len(_encoding.encode(piece)) <= 100

    def test_adjacent_pieces_actually_overlap(self):
        text = " ".join(f"word{i}" for i in range(500))
        pieces = _split_by_tokens(text, max_tokens=100, overlap_ratio=0.12)
        first_tail = _encoding.encode(pieces[0])[-5:]
        second_tokens = _encoding.encode(pieces[1])
        # the last few tokens of piece 1 should reappear near the start of piece 2
        assert any(second_tokens[i : i + 5] == first_tail for i in range(min(30, len(second_tokens) - 5)))


class TestChunkDocument:
    def test_chunk_ref_uses_heading_when_section_fits_in_one_chunk(self):
        loaded = LoadedDocument(format="text", text="# Eligibility\nShort section body.")
        chunks = chunk_document(loaded)
        assert len(chunks) == 1
        assert chunks[0].chunk_ref == "Eligibility"

    def test_chunk_ref_gets_part_suffix_when_section_splits(self):
        long_body = " ".join(f"word{i}" for i in range(2000))
        loaded = LoadedDocument(format="text", text=f"# Big Section\n{long_body}")
        chunks = chunk_document(loaded)
        assert len(chunks) > 1
        assert chunks[0].chunk_ref == f"Big Section (part 1/{len(chunks)})"
        assert chunks[-1].chunk_ref == f"Big Section (part {len(chunks)}/{len(chunks)})"


class TestLoadDocument:
    def test_docx_heading_styles_become_markdown_headings(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("Eligibility", style="Heading 1")
        doc.add_paragraph("Employees must work at least 20 hours/week.")
        doc.add_paragraph("Approval Process", style="Heading 2")
        doc.add_paragraph("Submit a request to your manager.")
        path = tmp_path / "sample.docx"
        doc.save(str(path))

        loaded = load_document(path)

        assert loaded.format == "docx"
        assert "# Eligibility" in loaded.text
        assert "## Approval Process" in loaded.text

    def test_docx_heading_flows_through_to_chunk_ref(self, tmp_path):
        # End-to-end version of retrieval-quality-note.md row #5: a DOCX "Eligibility"
        # heading should show up as a human-readable chunk_ref, not a raw chunk index.
        doc = DocxDocument()
        doc.add_paragraph("Eligibility", style="Heading 1")
        doc.add_paragraph("Employees must work at least 20 hours/week to qualify.")
        path = tmp_path / "sample.docx"
        doc.save(str(path))

        chunks = chunk_document(load_document(path))

        assert any(c.chunk_ref == "Eligibility" for c in chunks)

    def test_html_headings_become_markdown_headings(self, tmp_path):
        html = "<html><body><h2>Refund Policy</h2><p>Refunds within 30 days.</p></body></html>"
        path = tmp_path / "sample.html"
        path.write_text(html)

        loaded = load_document(path)

        assert loaded.format == "html"
        assert "## Refund Policy" in loaded.text
        assert "Refunds within 30 days." in loaded.text
